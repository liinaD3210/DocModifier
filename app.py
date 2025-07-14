import os
from dotenv import load_dotenv

load_dotenv() # Эта команда загружает переменные из .env файла

import streamlit as st
from docx import Document
from io import BytesIO
import json # Оставляем, так как может использоваться в get_diff_for_instruction или format_instruction_for_display
import html
# import textwrap # По-прежнему не вижу его использования, можно удалить, если уверены

try:
    from core.llm_handler import build_graph, GraphState # Убедитесь, что llm_handler содержит build_graph
    from core.docx_modifier import extract_text_from_doc, modify_document_with_structured_instructions
    # find_paragraphs_with_text не используется напрямую в этом app.py, убрал для чистоты, если не нужен
except ImportError as e:
    st.error(f"Критическая ошибка импорта: {e}. Убедитесь, что все файлы 'core' на месте и имена корректны.")
    st.stop()

EXAMPLE_DOC_PATH = "example_document.docx" # Убедитесь, что этот файл существует в корне проекта

# --- Конфигурация страницы ---
st.set_page_config(
    page_title="Агент правок DOCX",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Инициализация и кэширование графа ---
@st.cache_resource
def get_graph_instance(): # Переименовал для ясности
    try:
        graph = build_graph() # build_graph из вашего llm_handler.py (или llm_graph_builder.py)
        return graph
    except Exception as e:
        st.error(f"Не удалось инициализировать LangGraph: {e}")
        return None

if 'app_graph' not in st.session_state:
    st.session_state.app_graph = get_graph_instance()

# NEW_FEATURE_START: Функция загрузки примера
def load_example_document():
    if os.path.exists(EXAMPLE_DOC_PATH):
        with open(EXAMPLE_DOC_PATH, "rb") as f:
            return f.read(), os.path.basename(EXAMPLE_DOC_PATH)
    else:
        # Это сообщение будет видно, если пример не найден при попытке его загрузить
        st.toast(f"Файл примера '{EXAMPLE_DOC_PATH}' не найден. Функционал примера недоступен.", icon="⚠️")
        return None, None
# NEW_FEATURE_END

def init_session_state(clear_all=False, load_example_on_first_ever_run=False): # Изменил имя параметра
    graph_instance = st.session_state.get('app_graph')

    if clear_all:
        keys_to_preserve = {'app_graph'} # Сохраняем только граф
        preserved_values = {k: st.session_state[k] for k in keys_to_preserve if k in st.session_state}
        
        for key in list(st.session_state.keys()): # Очищаем все ключи
            del st.session_state[key]
        
        for k, v in preserved_values.items(): # Восстанавливаем сохраненные (граф)
            st.session_state[k] = v

    defaults = {
        "chat_messages": [], "current_doc_bytes": None, "original_file_name": None,
        "doc_loaded_flag": False, # Общий флаг, что какой-либо документ (пользовательский или пример) загружен
        "is_example_active": False, # Флаг, что активен именно пример
        "processing": False, "show_confirmation": False, 
        "proposed_instructions": None, "awaiting_clarification": False,
        "user_made_first_query_on_current_doc": False # Флаг для инструкции "Как пользоваться" для текущего документа
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value
    
    # NEW_FEATURE_START: Логика загрузки примера при первом запуске
    if load_example_on_first_ever_run and not st.session_state.doc_loaded_flag:
        example_bytes, example_name = load_example_document()
        if example_bytes:
            st.session_state.current_doc_bytes = example_bytes
            st.session_state.original_file_name = example_name
            st.session_state.doc_loaded_flag = True
            st.session_state.is_example_active = True
            st.session_state.user_made_first_query_on_current_doc = False # Для примера инструкция в чате
            st.session_state.chat_messages = [
                {"role": "assistant", 
                 "content": (f"Добро пожаловать! 👋 Я загрузил пример документа: **'{example_name}'**.\n\n"
                             f"Можете скачать текущую версию, внести изменения, после чего сравнить исходную и изменную версии примера.\n\n"
                             f"Вы можете попробовать следующие команды:\n"
                             f"* `Замени [Название Проекта] на 'Проект Аврора'`\n"
                             f"* `Удали абзац, который начинается с 'Этот абзац предназначен для демонстрации удаления.'`\n"
                             f"* `Сделай текст 'Важное замечание' жирным и подчеркнутым`\n"
                             f"* `Выровняй по центру абзац За неисполнение или ненадлежащее исполнение...`\n\n"
                             f"Или загрузите свой документ на панели слева. Удачи!")}
            ]
            # st.toast не нужен здесь, т.к. приветственное сообщение уже есть
    # NEW_FEATURE_END

# NEW_FEATURE_START: Условие для самого первого запуска сессии
# Проверяем наличие любого из наших ключей. Если ни одного нет, это самый первый запуск.
# Или если 'chat_messages' не существует (например, после st.experimental_rerun без сохранения всего состояния)
if not any(key in st.session_state for key in ["chat_messages", "doc_loaded_flag"]):
    init_session_state(load_example_on_first_ever_run=True)
# NEW_FEATURE_END

# --- Функции-обработчики ---
# ВАШИ ФУНКЦИИ get_diff_for_instruction, show_confirmation_ui, 
# handle_user_prompt, handle_user_confirmation ОСТАЮТСЯ ЗДЕСЬ БЕЗ ИЗМЕНЕНИЙ
# Я их скопирую из вашего предоставленного кода.

def get_diff_for_instruction(instruction: dict, doc: Document) -> dict:
    """
    ФИНАЛЬНАЯ ВЕРСИЯ: Готовит "было/стало" с HTML-выделением изменений и тусклым контекстом из слов.
    """
    result = {'before': 'Ошибка', 'after': 'Ошибка', 'notes': 'Не удалось обработать правку.', 'found': False}
    
    if not doc:
        result['notes'] = 'Объект документа не был передан.'
        return result

    try:
        op_type = instruction.get("operation_type")
        target = instruction.get("target_description", {})
        params = instruction.get("parameters", {})
        
        search_text = target.get("text_to_find")
        if not search_text:
            if op_type == "REPLACE_TEXT":
                search_text = params.get("old_text")
            elif op_type == "APPLY_FORMATTING":
                search_text = params.get("apply_to_text_segment")
        
        if not search_text:
            result['notes'] = 'LLM не предоставила достаточно данных для поиска.'
            return result

        # --- НАЧАЛО ВАШЕЙ ЛОГИКИ get_diff_for_instruction ---
        full_text_str = "\n".join([p.text for p in doc.paragraphs]) # Упрощенно, лучше через extract_text_from_doc
        all_words = full_text_str.split()
        search_words = search_text.split()
        target_word_start_index = -1
        for i in range(len(all_words) - len(search_words) + 1):
            if all_words[i:i+len(search_words)] == search_words:
                target_word_start_index = i
                break

        if target_word_start_index == -1:
            result['notes'] = f'Текст «{html.escape(search_text)}» не был найден для предпросмотра.'
            return result
        
        target_word_end_index = target_word_start_index + len(search_words)
        context_words_count = 10 # Уменьшил для краткости в UI
        start_idx = max(0, target_word_start_index - context_words_count)
        end_idx = min(len(all_words), target_word_end_index + context_words_count)
        
        words_before_context = all_words[start_idx:target_word_start_index]
        words_of_target = all_words[target_word_start_index:target_word_end_index]
        words_after_context = all_words[target_word_end_index:end_idx]
        
        style_context = "opacity: 0.6;"
        style_highlight_before = "background-color: #FFD2D2; color: #A62020; padding: 1px 3px; border-radius: 3px; font-weight: bold;"
        style_highlight_after = "background-color: #D2FFD2; color: #206620; padding: 1px 3px; border-radius: 3px; font-weight: bold;"
        style_highlight_format = "background-color: #D0E0FF; color: #103050; padding: 1px 3px; border-radius: 3px; font-style: italic;" # Добавил курсив для наглядности
        
        escaped_context_before = html.escape(" ".join(words_before_context))
        escaped_target = html.escape(" ".join(words_of_target))
        escaped_context_after = html.escape(" ".join(words_after_context))

        result['before'] = (
            f"<span style='{style_context}'>...{escaped_context_before}</span> "
            f"<span style='{style_highlight_before}'>{escaped_target}</span> "
            f"<span style='{style_context}'>{escaped_context_after}...</span>"
        )
        
        notes = f"Операция: `{op_type}`. "
        after_html = result['before'] # По умолчанию, если операция не меняет текст напрямую

        if op_type == "REPLACE_TEXT":
            old, new = params.get("old_text", search_text), params.get("new_text", "")
            escaped_new = html.escape(new)
            # Заменяем только целевую часть, оставляя контекст как был "до"
            after_html = (
                f"<span style='{style_context}'>...{escaped_context_before}</span> "
                f"<span style='{style_highlight_after}'>{escaped_new}</span> " # Показываем новый текст вместо старого
                f"<span style='{style_context}'>{escaped_context_after}...</span>"
            )
            notes += f"Замена «{html.escape(old)}» на «{html.escape(new)}»."
        elif op_type == "INSERT_TEXT":
            to_insert = params.get("text_to_insert", "")
            position = params.get("position", "after_paragraph") # Уточнить позицию для отображения
            escaped_insert = html.escape(to_insert)
            if "after" in position: # Упрощенное отображение для вставки
                after_html = (
                    f"<span style='{style_context}'>...{escaped_context_before}</span> "
                    f"{escaped_target} <span style='{style_highlight_after}'>{escaped_insert}</span> "
                    f"<span style='{style_context}'>{escaped_context_after}...</span>"
                )
            elif "before" in position:
                 after_html = (
                    f"<span style='{style_context}'>...{escaped_context_before}</span> "
                    f"<span style='{style_highlight_after}'>{escaped_insert}</span> {escaped_target} "
                    f"<span style='{style_context}'>{escaped_context_after}...</span>"
                )
            else: # start_of_paragraph, end_of_paragraph - сложнее точно показать в этом diff
                after_html = f"{result['before']} <span style='{style_highlight_after}'>(вставлено: {escaped_insert})</span>"

            notes += f"Вставка текста: «{escaped_insert}» ({position})."
        elif op_type == "DELETE_ELEMENT":
            after_html = (
                 f"<span style='{style_context}'>...{escaped_context_before}</span> "
                 f"<span style='text-decoration: line-through; color: #FFAAAA; background-color: #502020;'> (удаленный элемент) </span> "
                 f"<span style='{style_context}'>{escaped_context_after}...</span>"
            )
            notes += f"Удаление элемента, содержащего «{escaped_target}»."
        elif op_type == "APPLY_FORMATTING":
            # Для форматирования, "стало" будет выглядеть так же, но с примененным стилем
            # Мы не можем показать это в HTML без реального применения и сравнения XML или сложного рендеринга.
            # Поэтому просто опишем действие.
            rules_str_list = []
            for r_item in params.get("formatting_rules", []): # Переименовал переменную цикла
                rules_str_list.append(f"`{r_item.get('style')}`: `{r_item.get('value')}`")
            applied_formatting_desc = ", ".join(rules_str_list)
            
            after_html = ( # Показываем целевой текст с подсветкой форматирования
                f"<span style='{style_context}'>...{escaped_context_before}</span> "
                f"<span style='{style_highlight_format}'>{escaped_target}</span> " # Целевой текст выделен
                f"<span style='{style_context}'>{escaped_context_after}...</span>"
            )
            notes += f"Будет применено форматирование к «{escaped_target}»: {applied_formatting_desc}."
        
        result['after'] = after_html
        result['notes'] = notes
        result['found'] = True
        # --- КОНЕЦ ВАШЕЙ ЛОГИКИ get_diff_for_instruction ---
    except Exception as e:
        result['notes'] = f"Ошибка при генерации предпросмотра: {e}"
    return result

def show_confirmation_ui(instructions: list[dict]):
    if "selected_instructions" not in st.session_state:
        st.session_state.selected_instructions = {i: True for i in range(len(instructions))}

    st.subheader("🤖 Проверьте и подтвердите правки")
    st.caption("Снимите галочки с правок, которые вы не хотите применять.")
    st.markdown("---")
    
    doc_object_for_diff = None # Инициализируем
    if st.session_state.current_doc_bytes:
        try:
            doc_object_for_diff = Document(BytesIO(st.session_state.current_doc_bytes))
        except Exception as e:
            st.warning(f"Не удалось загрузить документ для предпросмотра diff: {e}")
            doc_object_for_diff = None # Убедимся, что None, если ошибка

    container_style = "padding: 0.5rem; border: 1px solid #4A4A4A; border-radius: 0.3rem; margin-bottom: 0.5rem; background-color: #262730; color: #FAFAFA;"
    notes_style = "font-size: 0.9em; color: #A0A0A0;"


    for i, instruction in enumerate(instructions):
        with st.container(border=True): # Внешний контейнер для каждой правки
            op_type = instruction.get("operation_type", "Неизвестная операция")
            
            cols_header = st.columns([0.05, 0.95])
            with cols_header[0]:
                is_selected = st.checkbox(" ", value=st.session_state.selected_instructions.get(i, True), key=f"cb_diff_{i}", label_visibility="collapsed")
                st.session_state.selected_instructions[i] = is_selected
            with cols_header[1]:
                st.markdown(f"##### Правка {i+1}: `{op_type}`")

            if doc_object_for_diff: # Только если документ успешно загружен для diff
                diff = get_diff_for_instruction(instruction, doc_object_for_diff)
                
                if diff['found']:
                    st.markdown("**Было (контекст):**")
                    st.markdown(f"<div style='{container_style}'>{diff['before']}</div>", unsafe_allow_html=True)
                    
                    st.markdown("**Станет (контекст):**")
                    st.markdown(f"<div style='{container_style}'>{diff['after']}</div>", unsafe_allow_html=True)
                else: # Если текст для diff не найден
                    st.markdown(f"**Описание действия:** {format_instruction_for_display(instruction)}")


                if diff['notes']:
                    st.markdown(f"<div style='{notes_style}'>ℹ️ {html.escape(diff['notes'])}</div>", unsafe_allow_html=True)
            else: # Если doc_object_for_diff is None
                st.markdown(f"**Описание действия:** {format_instruction_for_display(instruction)}")
                st.caption("Предпросмотр изменений недоступен, так как не удалось обработать текущий документ.")

        st.markdown("<br>", unsafe_allow_html=True) 

    st.markdown("---")
    apply_col, cancel_col, _ = st.columns([2, 1, 3])
    if apply_col.button("✅ Применить выбранные правки", use_container_width=True, key="apply_btn_confirmation"):
        handle_user_confirmation(approved=True)
    if cancel_col.button("❌ Отклонить все", use_container_width=True, key="cancel_btn_confirmation"):
        handle_user_confirmation(approved=False)

def handle_user_prompt(user_input: str):
    st.session_state.processing = True
    st.session_state.chat_messages.append({"role": "user", "content": user_input})
    try:
        if not st.session_state.current_doc_bytes:
            st.error("Документ не загружен. Пожалуйста, загрузите документ перед отправкой запроса.")
            st.session_state.chat_messages.append({"role": "assistant", "content": "Ошибка: Документ не загружен."})
            st.session_state.processing = False # Сбрасываем флаг
            st.rerun() # Перерисовываем, чтобы показать ошибку
            return # Выходим из функции

        doc_content = extract_text_from_doc(Document(BytesIO(st.session_state.current_doc_bytes)))
        initial_state = GraphState(
            original_user_query=user_input, current_user_query=user_input,
            document_content_text=doc_content, document_bytes=st.session_state.current_doc_bytes,
            extracted_instructions=None, clarification_question=None, system_message=None, next_node_to_call=None
        )
        with st.spinner("🤖 Агент анализирует ваш запрос..."):
            if not st.session_state.app_graph: # Дополнительная проверка
                st.error("Критическая ошибка: Граф обработки не инициализирован.")
                st.session_state.chat_messages.append({"role": "assistant", "content": "Ошибка конфигурации агента. Попробуйте перезагрузить страницу."})
                st.session_state.processing = False
                st.rerun()
                return

            final_state = st.session_state.app_graph.invoke(initial_state, {"recursion_limit": 15})

        st.session_state.awaiting_clarification = bool(final_state.get("clarification_question"))
        if final_state.get("extracted_instructions"):
            st.session_state.proposed_instructions = final_state["extracted_instructions"]
            st.session_state.show_confirmation = True
        elif final_state.get("clarification_question"):
            st.session_state.chat_messages.append({"role": "assistant", "content": final_state["clarification_question"]})
        elif final_state.get("system_message"):
            st.session_state.chat_messages.append({"role": "assistant", "content": final_state["system_message"]})
        else:
             st.session_state.chat_messages.append({"role": "assistant", "content": "Не удалось выработать план действий. Попробуйте переформулировать."})
    except Exception as e:
        st.error(f"Ошибка при анализе запроса: {e}")
        st.session_state.chat_messages.append({"role": "assistant", "content": f"Произошла критическая ошибка: {e}"})
    finally:
        st.session_state.processing = False
        st.rerun()

def handle_user_confirmation(approved: bool):
    if not approved:
        st.session_state.chat_messages.append({"role": "assistant", "content": "Предложенные действия были отклонены."})
    else:
        # Используем .get() для selected_instructions для безопасного доступа
        selected_instructions_map = st.session_state.get("selected_instructions", {})
        selected_indices = [i for i, sel in selected_instructions_map.items() if sel]
        
        instructions_to_apply = []
        if st.session_state.proposed_instructions: # Проверяем, что список существует
            instructions_to_apply = [
                st.session_state.proposed_instructions[i] 
                for i in selected_indices 
                if i < len(st.session_state.proposed_instructions) # Доп. проверка на выход за пределы
            ]

        if not instructions_to_apply:
            st.session_state.chat_messages.append({"role": "assistant", "content": "Вы не выбрали ни одной правки для применения. Действия отменены."})
        else:
            st.session_state.processing = True
            st.session_state.chat_messages.append({"role": "assistant", "content": f"Применяю {len(instructions_to_apply)} подтвержденных изменений..."})
            try:
                if not st.session_state.current_doc_bytes: # Проверка
                    st.error("Документ не загружен. Невозможно применить изменения.")
                    st.session_state.chat_messages.append({"role": "assistant", "content": "Ошибка: Документ для применения правок не найден."})
                    return # Выходим из функции

                doc = Document(BytesIO(st.session_state.current_doc_bytes))
                success = modify_document_with_structured_instructions(doc, instructions_to_apply)
                if success:
                    bio = BytesIO()
                    doc.save(bio)
                    st.session_state.current_doc_bytes = bio.getvalue()
                    st.session_state.chat_messages.append({"role": "assistant", "content": "Изменения успешно применены."})
                else:
                    st.session_state.chat_messages.append({"role": "assistant", "content": "Не удалось применить некоторые или все изменения (возможно, текст не найден или произошла ошибка в обработчике)."})
            except Exception as e:
                st.error(f"Ошибка при применении изменений: {e}")
                st.session_state.chat_messages.append({"role": "assistant", "content": f"Ошибка выполнения: {e}"})
            finally:
                st.session_state.processing = False # Сбрасываем флаг здесь
    
    # Сброс состояния подтверждения происходит в любом случае (approved или not)
    st.session_state.show_confirmation = False
    st.session_state.proposed_instructions = None
    if "selected_instructions" in st.session_state: # Безопасное удаление
        del st.session_state.selected_instructions
    
    # st.rerun() вызывается после finally в handle_user_prompt, или здесь, если нужно обновить UI немедленно
    # Если processing был True, то rerun из handle_user_prompt может не случиться, если была ошибка.
    # Лучше иметь rerun здесь, чтобы UI обновился после этого действия.
    if not st.session_state.processing: # Только если не идем в processing в handle_user_prompt
        st.rerun()


def format_instruction_for_display(instruction: dict) -> str:
    op_type = instruction.get("operation_type", "Неизвестная операция")
    params = instruction.get("parameters", {})
    target = instruction.get("target_description", {})
    
    display_parts = [f"**Действие:** `{op_type}`"]
    
    if op_type == "REPLACE_TEXT":
        old = params.get('old_text', 'N/A')
        new = params.get('new_text', 'N/A')
        context = target.get('text_to_find')
        display_parts.append(f"- Заменить: `{html.escape(old)}`")
        display_parts.append(f"- На: `{html.escape(new)}`")
        if context: display_parts.append(f"- В контексте: `{html.escape(context)}`")
    elif op_type == "INSERT_TEXT":
        text_ins = params.get('text_to_insert', 'N/A')
        pos = params.get('position', 'N/A')
        context = target.get('text_to_find')
        display_parts.append(f"- Вставить: `{html.escape(text_ins)}`")
        display_parts.append(f"- Позиция: `{pos}`")
        if context: display_parts.append(f"- Относительно: `{html.escape(context)}`")
    elif op_type == "DELETE_ELEMENT":
        el_type = target.get('element_type', 'N/A')
        context = target.get('text_to_find')
        display_parts.append(f"- Удалить элемент типа: `{el_type}`")
        if context: display_parts.append(f"- Идентифицированный по тексту: `{html.escape(context)}`")
    elif op_type == "APPLY_FORMATTING":
        rules_display = [f"  - `{r.get('style')}`: `{r.get('value')}`" for r in params.get("formatting_rules", [])]
        context = target.get('text_to_find')
        segment = params.get('apply_to_text_segment')
        target_display = segment if segment else context
        if target_display:
            display_parts.append(f"- Применить форматирование к: `{html.escape(target_display)}`")
        if rules_display: display_parts.extend(rules_display)
    else:
        if params: display_parts.append(f"- Параметры: ```json\n{json.dumps(params, indent=2, ensure_ascii=False)}\n```")
        if target: display_parts.append(f"- Цель: ```json\n{json.dumps(target, indent=2, ensure_ascii=False)}\n```")
    return "\n".join(display_parts)
# --- КОНЕЦ ВАШИХ ФУНКЦИЙ ---

# --- Основной UI ---
st.title("📄 Агент правок DOCX") # Упростил заголовок
st.caption("Загрузите документ, опишите правки текстом, просмотрите и подтвердите изменения.")

with st.sidebar:
    st.header("Управление документом")
    
    # NEW_FEATURE_START: Логика отображения кнопки "Попробовать с примером"
    if not st.session_state.doc_loaded_flag: # Если никакой документ еще не был загружен (ни пример, ни пользовательский)
        if st.button("🚀 Попробовать с примером", use_container_width=True, key="load_example_sidebar_btn"):
            init_session_state(clear_all=True, load_example_on_first_ever_run=True) # Перезагружаем сессию с примером
            st.rerun()
    # NEW_FEATURE_END

    uploaded_file_widget = st.file_uploader( # Даем явное имя виджету
        "Или загрузите свой .docx файл:", 
        type=["docx"], 
        key="user_doc_uploader_main", # Изменил ключ, чтобы избежать конфликтов
        disabled=st.session_state.processing
    )
    
    # NEW_FEATURE_START: Логика обработки загрузки пользовательского файла
    if uploaded_file_widget and \
       (uploaded_file_widget.name != st.session_state.original_file_name or not st.session_state.doc_loaded_flag): # Если загружен новый файл или до этого ничего не было
        init_session_state(clear_all=True) # Сбрасываем все, включая флаги примера
        st.session_state.current_doc_bytes = uploaded_file_widget.getvalue()
        st.session_state.original_file_name = uploaded_file_widget.name
        st.session_state.doc_loaded_flag = True
        st.session_state.is_example_active = False # Явно указываем, что это не пример
        st.session_state.user_made_first_query_on_current_doc = False # Сбрасываем для нового документа
        st.session_state.chat_messages = [ # Чистим чат и добавляем сообщение о загрузке
            {"role": "assistant", "content": f"Файл **'{uploaded_file_widget.name}'** успешно загружен. Готов к вашим командам!"}
        ]
        st.toast(f"Файл '{uploaded_file_widget.name}' загружен.", icon="👍")
        st.rerun() # Важно для обновления UI после загрузки
    # NEW_FEATURE_END

    if st.session_state.doc_loaded_flag: # Если какой-либо документ загружен
        st.info(f"Активный документ: **{st.session_state.original_file_name}**")
        if st.session_state.is_example_active:
            st.caption("Это демонстрационный документ.")

        # NEW_FEATURE_START: Кнопка сброса теперь всегда сбрасывает к состоянию "ожидание загрузки" или к примеру
        if st.button("Загрузить другой / Сбросить", use_container_width=True, key="reset_sidebar_main_btn",
                      disabled=st.session_state.processing):
            init_session_state(clear_all=True, load_example_on_first_ever_run=True) # При сбросе снова пытаемся загрузить пример
            st.rerun()
        # NEW_FEATURE_END
        
        if st.session_state.current_doc_bytes:
            download_file_name = f"{'example_modified' if st.session_state.is_example_active else 'modified'}_{st.session_state.original_file_name or 'document.docx'}"
            st.download_button("⬇️ Скачать текущий документ", st.session_state.current_doc_bytes,
                download_file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, disabled=st.session_state.processing, key="final_download_main_btn"
            )
    else: 
        st.caption("Загрузите свой .docx или попробуйте с примером.")

    st.divider()
    st.caption("**Proof of Concept (v0.1)**") # Используем st.caption для заголовка
    st.caption("""
    Это демонстрационная версия, подтверждающая основную концепцию. Проект открыт для дальнейших доработок и улучшения.
    """)
    st.caption("""
    Мы не храним загруженные файлы.
    """)


if not st.session_state.doc_loaded_flag: # Если никакой документ не загружен
    st.info("👈 Пожалуйста, загрузите .docx документ или **попробуйте с примером** на панели слева, чтобы начать.")
elif not st.session_state.app_graph:
    st.error("Ошибка инициализации агента (LangGraph). Функционал недоступен.")
else:
    # NEW_FEATURE_START: Условие для инструкции "Как пользоваться"
    # Показываем, если загружен НЕ пример И пользователь еще не делал запросов к ЭТОМУ документу
    if not st.session_state.user_made_first_query_on_current_doc and not st.session_state.is_example_active:
        with st.container(border=True):
             st.subheader("💡 Как начать:")
             st.markdown("""
             1. Убедитесь, что ваш `.docx` документ **загружен**.
             2. **Опишите правки** в поле ввода ниже (например: "Замени 'Старый Текст' на 'Новый Текст' в первом абзаце").
             3. Система предложит изменения. **Просмотрите** их внимательно.
             4. **Подтвердите или отклоните** предложенные правки.
             5. При необходимости, **скачайте** обновленный документ или **продолжите вносить правки**.
             """)
        st.markdown("---")
    # NEW_FEATURE_END

    # Отображение чата
    for msg in st.session_state.chat_messages: # Используем новый ключ
        with st.chat_message(msg["role"]):
            if isinstance(msg["content"], list):
                for item_md in msg["content"]:
                    st.markdown(item_md)
            else:
                st.markdown(msg["content"])

    if st.session_state.show_confirmation and st.session_state.proposed_instructions:
        show_confirmation_ui(st.session_state.proposed_instructions)
    
    chat_input_disabled_reason = None
    if st.session_state.processing: chat_input_disabled_reason = "Идет обработка..."
    elif st.session_state.show_confirmation: chat_input_disabled_reason = "Ожидание подтверждения действий..."
    
    prompt_for_chat_input = "Пожалуйста, ответьте на уточняющий вопрос:" if st.session_state.awaiting_clarification else "Что бы вы хотели изменить в документе?"
    
    if user_input_str := st.chat_input(prompt_for_chat_input, disabled=bool(chat_input_disabled_reason), key="main_chat_input_field_key"): # Изменил ключ
        if not st.session_state.user_made_first_query_on_current_doc: 
            st.session_state.user_made_first_query_on_current_doc = True # Отмечаем, что первый запрос для ТЕКУЩЕГО документа сделан
        handle_user_prompt(user_input_str)

    if chat_input_disabled_reason:
        st.caption(f"_{chat_input_disabled_reason}_")