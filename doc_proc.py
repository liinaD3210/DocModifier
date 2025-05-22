import os
import json
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_COLOR_INDEX # <--- ДОБАВЛЕН ИМПОРТ

import google.generativeai as genai

# Загрузка переменных окружения из .env файла
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    print("Ошибка: GOOGLE_API_KEY не найден в .env файле.")

if GOOGLE_API_KEY:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
    except Exception as e:
        print(f"Ошибка при конфигурации Gemini API: {e}")
        GOOGLE_API_KEY = None

# --- Функция для LLM с использованием Gemini (для нескольких правок) ---
def get_llm_instructions_list(doc_content_text: str, user_query: str) -> list[dict] | None:
    """
    Обращается к Google Gemini API для получения списка инструкций по замене текста.

    Args:
        doc_content_text (str): Текстовое содержимое документа.
        user_query (str): Запрос пользователя, который может содержать несколько правок.

    Returns:
        list[dict] | None: Список словарей, где каждый словарь имеет ключи 
                             "old_text" и "new_text", или None, если LLM не смогла помочь.
    """
    if not GOOGLE_API_KEY:
        print("ПРЕДУПРЕЖДЕНИЕ: Работа в режиме заглушки LLM (множественные правки).")
        # Заглушка для нескольких правок
        instructions = []
        if "цену договора на 10 000 000 рублей" in user_query.lower():
            # Предположим, в документе есть "Цена Договора: 5 000 000 (Пять миллионов) рублей 00 копеек."
            # LLM должна быть достаточно умной, чтобы найти "5 000 000" или "5 000 000 (Пять миллионов) рублей 00 копеек"
            # Для заглушки упростим:
            if "5 000 000" in doc_content_text:
                 instructions.append({"old_text": "5 000 000", "new_text": "10 000 000"})
        if "предоплату на 75%" in user_query.lower():
            if "50%" in doc_content_text: # Предположим, в документе "Предоплата составляет 50% от Цены Договора."
                 instructions.append({"old_text": "50%", "new_text": "75%"})
        
        if "дату договора на 24.04.2025" in user_query.lower():
             if "15.03.2024" in doc_content_text:
                 instructions.append({"old_text": "15.03.2024", "new_text": "24.04.2025"})
             elif "[ДАТА_ДОГОВОРА]" in doc_content_text:
                 instructions.append({"old_text": "[ДАТА_ДОГОВОРА]", "new_text": "24.04.2025"})
        
        return instructions if instructions else None

    model_name = "gemini-1.5-flash-latest"
    generation_config = {"response_mime_type": "application/json"}
    model = genai.GenerativeModel(model_name, generation_config=generation_config)

    prompt = f"""
Тебе предоставлен текст документа и запрос пользователя на внесение одного или НЕСКОЛЬКИХ изменений.
Твоя задача - для КАЖДОГО запрошенного изменения точно определить:
1. Фрагмент текста в документе, который нужно заменить (`old_text`).
2. Текст, на который его нужно заменить (`new_text`).

Очень важно:
- `old_text` должен быть ТОЧНЫМ фрагментом из документа, включая регистр, знаки препинания и пробелы.
- Если для какого-то конкретного изменения ты не можешь точно определить `old_text` или `new_text`, пропусти это изменение, но обработай остальные.
- Если не найдено ни одного изменения, верни пустой список.

Текст документа (фрагмент для анализа):
---
{doc_content_text[:15000]} 
---
Запрос пользователя: "{user_query}"

Верни результат строго в формате JSON-массива объектов. Каждый объект должен содержать ключи "old_text" и "new_text".
Если изменение невозможно определить, `old_text` или `new_text` могут быть `null` для этого конкретного изменения, или объект может быть пропущен.
Предпочтительно пропускать объект, если `old_text` не найден.

Пример формата ответа для нескольких изменений:
[
  {{"old_text": "точный текст первой замены", "new_text": "новый текст первой замены"}},
  {{"old_text": "точный текст второй замены", "new_text": "новый текст второй замены"}}
]

Пример формата ответа для одного изменения:
[
  {{"old_text": "существующий текст", "new_text": "текст для замены"}}
]

Пример формата ответа, если ничего не найдено или не понято:
[]

Проанализируй текст документа и запрос пользователя, и предоставь JSON-массив.
Убедись, что `old_text` включает всю пунктуацию, как в оригинальном тексте, включая точки, запятые и т.д. в конце фрагмента, если они там есть.
"""
    # Увеличил немного контекст до 15000 символов

    print(f"\n--- Отправка запроса в Gemini API ({model_name}) для нескольких правок ---")
    print(f"Запрос пользователя: {user_query}")
    
    try:
        response = model.generate_content(prompt)
        
        print("--- Ответ от Gemini API (сырой текст) ---")
        print(response.text) 

        # Gemini в JSON-режиме должен возвращать валидный JSON
        parsed_response = json.loads(response.text)
        
        print("--- Ответ от Gemini API (распарсенный JSON) ---")
        print(parsed_response)

        if isinstance(parsed_response, list):
            valid_instructions = []
            for item in parsed_response:
                if isinstance(item, dict) and \
                   item.get("old_text") is not None and \
                   item.get("new_text") is not None and \
                   item.get("old_text") != "": # old_text не должен быть пустым
                    valid_instructions.append({
                        "old_text": str(item["old_text"]),
                        "new_text": str(item["new_text"])
                    })
                else:
                    print(f"ПРЕДУПРЕЖДЕНИЕ: Пропущен некорректный элемент от LLM: {item}")
            return valid_instructions if valid_instructions else None
        else:
            print(f"ПРЕДУПРЕЖДЕНИЕ: LLM вернула не список, а {type(parsed_response)}.")
            return None

    except json.JSONDecodeError as e:
        print(f"Ошибка декодирования JSON ответа от Gemini: {e}")
        print(f"Полученный текст: {response.text if 'response' in locals() else 'Ответ не получен'}")
        return None
    except Exception as e:
        print(f"Ошибка при взаимодействии с Gemini API: {e}")
        if hasattr(e, 'response') and hasattr(e.response, 'prompt_feedback'):
            print(f"Prompt Feedback: {e.response.prompt_feedback}")
        return None

# --- Функции замены текста в DOCX (остаются прежними, т.к. modify_docx вызывается для каждой правки) ---

def _replace_text_in_paragraph_runs_with_highlight(p, old_text, new_text): # Переименовал для ясности, что она с подсветкой
    modified_paragraph = False
    if not old_text or old_text not in p.text:
        return False

    runs = p.runs
    # Простой случай: old_text целиком в одном run'е
    # Этот блок можно улучшить, чтобы он не конфликтовал с логикой сложной замены,
    # но для начала оставим его для простых и быстрых замен.
    for i, run in enumerate(runs):
        if old_text in run.text:
            # Если old_text - это весь текст run'а или его часть, но он не разбит
            # (эвристика: если нет другого вхождения в полном тексте параграфа, которое начинается раньше)
            # Простая проверка:
            if p.text.find(old_text) == (sum(len(r.text) for r in runs[:i]) + run.text.find(old_text)):
                print(f"DEBUG (простая замена): old_text ('{old_text}') НАЙДЕН в одном run: '{run.text}'")
                # Заменяем только первое вхождение в этом run, чтобы избежать проблем, если new_text содержит old_text
                current_run_text = run.text
                start_replace_index = current_run_text.find(old_text)
                end_replace_index = start_replace_index + len(old_text)
                run.text = current_run_text[:start_replace_index] + new_text + current_run_text[end_replace_index:]
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                return True # Изменение сделано, выходим из этой функции для данного параграфа

    # Сложный случай: old_text разбит на несколько run'ов
    current_pos = 0
    text_segments = [] 
    para_full_text = ""
    for i, run in enumerate(runs):
        text_segments.append((i, current_pos, run.text))
        para_full_text += run.text
        current_pos += len(run.text)

    start_match_idx = para_full_text.find(old_text)
    if start_match_idx != -1:
        print(f"DEBUG (сложная замена): old_text ('{old_text}') НАЙДЕН в склеенном тексте параграфа.")
        end_match_idx = start_match_idx + len(old_text)

        first_run_idx_involved = -1
        
        # Определяем, какие run'ы задействованы
        # Ищем первый run, который содержит или начинает old_text
        for run_idx, run_start_pos, run_text_content in text_segments:
            run_end_pos = run_start_pos + len(run_text_content)
            if run_start_pos <= start_match_idx < run_end_pos:
                first_run_idx_involved = run_idx
                break
        
        if first_run_idx_involved != -1:
            # Собираем текст, начиная с первого задействованного run'а
            accumulated_text = ""
            runs_actually_involved_indices = [] # run'ы, которые действительно составляют old_text

            for run_idx_iter in range(first_run_idx_involved, len(runs)):
                runs_actually_involved_indices.append(run_idx_iter)
                accumulated_text += runs[run_idx_iter].text
                if len(accumulated_text) >= (end_match_idx - text_segments[first_run_idx_involved][1]): # Собрали достаточно текста
                    break
            
            # Проверка, что собранный текст действительно содержит old_text в начале (с учетом смещения)
            offset_in_first_involved_run = start_match_idx - text_segments[first_run_idx_involved][1]
            effective_accumulated_text = accumulated_text[offset_in_first_involved_run:]

            if effective_accumulated_text.startswith(old_text):
                # Модифицируем первый задействованный run
                first_run_obj = runs[runs_actually_involved_indices[0]]
                prefix_in_first_run = first_run_obj.text[:offset_in_first_involved_run]
                first_run_obj.text = prefix_in_first_run + new_text
                first_run_obj.font.highlight_color = WD_COLOR_INDEX.YELLOW
                modified_paragraph = True
                print(f"DEBUG: Заменен текст в Run {runs_actually_involved_indices[0]}. Новый префикс + new_text: '{first_run_obj.text}'")

                # Обрабатываем остальные задействованные run'ы
                remaining_old_text_len = len(old_text) - (len(runs[runs_actually_involved_indices[0]].text) - offset_in_first_involved_run - len(new_text) ) # сколько от old_text было в первом run после префикса
                
                # Эта часть сложная, нужно правильно рассчитать, сколько текста УДАЛИТЬ из последующих run'ов
                # Переписываем логику очистки последующих run'ов
                
                # Сколько от old_text было в первом run (за вычетом префикса)
                consumed_len_in_first_run = len(runs[runs_actually_involved_indices[0]].text) - len(prefix_in_first_run) - len(new_text)
                if consumed_len_in_first_run < 0: consumed_len_in_first_run = 0 # если new_text длиннее части old_text в первом run

                # Оставшаяся длина old_text, которую нужно "удалить" из последующих run'ов
                old_text_to_remove_len = len(old_text) - consumed_len_in_first_run

                for k in range(1, len(runs_actually_involved_indices)):
                    current_run_to_clear_idx = runs_actually_involved_indices[k]
                    run_to_clear = runs[current_run_to_clear_idx]
                    
                    if old_text_to_remove_len <= 0: break # Весь old_text уже "удален"

                    if len(run_to_clear.text) <= old_text_to_remove_len:
                        print(f"DEBUG: Очищается полностью Run {current_run_to_clear_idx}: старый текст '{run_to_clear.text}'")
                        old_text_to_remove_len -= len(run_to_clear.text)
                        run_to_clear.text = ""
                    else:
                        print(f"DEBUG: Очищается частично Run {current_run_to_clear_idx}: старый текст '{run_to_clear.text}', удаляем {old_text_to_remove_len} символов с начала")
                        run_to_clear.text = run_to_clear.text[old_text_to_remove_len:]
                        old_text_to_remove_len = 0 # Все удалено
                return True

        return modified_paragraph

# Остальные _replace_text_in_element_runs, process_document_elements, modify_docx, extract_text_from_doc
# остаются без изменений, так как `modify_docx` будет вызываться для каждой правки из списка.
# Важно, чтобы `_replace_text_in_paragraph_runs_with_highlight` корректно возвращала True/False.

def _replace_text_in_element_runs(element, old_text, new_text):
    modified_in_element = False
    if hasattr(element, 'paragraphs'):
        for p_in_el in element.paragraphs:
            if _replace_text_in_paragraph_runs_with_highlight(p_in_el, old_text, new_text):
                modified_in_element = True
    elif hasattr(element, 'runs'): # Для параграфов (когда element это Paragraph)
        if _replace_text_in_paragraph_runs_with_highlight(element, old_text, new_text):
            modified_in_element = True
    return modified_in_element

def process_document_elements(elements, old_text, new_text):
    modified_overall = False
    for el in elements:
        if _replace_text_in_element_runs(el, old_text, new_text):
            modified_overall = True
            # Если мы хотим, чтобы КАЖДОЕ вхождение old_text было заменено, то break здесь не нужен.
            # Если только первое вхождение в наборе elements, то break.
            # Для нашего случая (применение одной правки old->new) - break не нужен,
            # т.к. одна правка может затрагивать несколько мест в документе (если old_text не уникален).
            # Однако, если LLM дает уникальный old_text, то можно и break для оптимизации.
            # Пока оставим без break, чтобы заменить все вхождения текущей пары old/new.
    return modified_overall

def modify_docx(doc_object: Document, old_text: str, new_text: str) -> bool:
    overall_modified_for_this_edit = False 

    if process_document_elements(doc_object.paragraphs, old_text, new_text):
        overall_modified_for_this_edit = True

    for table in doc_object.tables:
        for row in table.rows:
            for cell in row.cells:
                if process_document_elements(cell.paragraphs, old_text, new_text):
                    overall_modified_for_this_edit = True
    
    for section in doc_object.sections:
        if process_document_elements(section.header.paragraphs, old_text, new_text):
            overall_modified_for_this_edit = True
        for table_in_header in section.header.tables:
            for row in table_in_header.rows:
                for cell in row.cells:
                    if process_document_elements(cell.paragraphs, old_text, new_text):
                        overall_modified_for_this_edit = True
        
        if process_document_elements(section.footer.paragraphs, old_text, new_text):
            overall_modified_for_this_edit = True
        for table_in_footer in section.footer.tables:
            for row in table_in_footer.rows:
                for cell in row.cells:
                    if process_document_elements(cell.paragraphs, old_text, new_text):
                        overall_modified_for_this_edit = True
        
    if overall_modified_for_this_edit:
        print(f"DEBUG: Правка ('{old_text}' -> '{new_text}') БЫЛА применена к документу.")
    else:
        print(f"DEBUG: Правка ('{old_text}' -> '{new_text}') НЕ была применена (текст не найден).")
    return overall_modified_for_this_edit

def extract_text_from_doc(doc_object: Document) -> str:
    full_text_parts = []
    for p in doc_object.paragraphs:
        full_text_parts.append(p.text)
    for table in doc_object.tables:
        for row in table.rows:
            for cell in row.cells:
                for p_in_cell in cell.paragraphs:
                    full_text_parts.append(p_in_cell.text)
    for section in doc_object.sections:
        for p_in_header in section.header.paragraphs:
            full_text_parts.append(p_in_header.text)
        for table_in_header in section.header.tables:
            for row in table_in_header.rows:
                for cell in row.cells:
                    for p_in_cell in cell.paragraphs:
                        full_text_parts.append(p_in_cell.text)
        for p_in_footer in section.footer.paragraphs:
            full_text_parts.append(p_in_footer.text)
        for table_in_footer in section.footer.tables:
            for row in table_in_footer.rows:
                for cell in row.cells:
                    for p_in_cell in cell.paragraphs:
                        full_text_parts.append(p_in_cell.text)
    return "\n".join(full_text_parts)