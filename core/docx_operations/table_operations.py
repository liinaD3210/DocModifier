# core/operations/table_operations.py
from docx import Document
from docx.table import Table, _Cell as CellType
from loguru import logger
from ..docx_utils import get_table_by_description # Относительный импорт

def handle_table_modify_cell(doc: Document, target_description: dict, parameters: dict) -> bool:
    # ... (ваш существующий код _handle_table_modify_cell, использующий get_table_by_description) ...
    logger.info(f"Выполнение TABLE_MODIFY_CELL: target={target_description}, params={parameters}")
    table_coords = target_description.get("table_coords")
    new_cell_text = parameters.get("new_cell_text")

    if not table_coords or new_cell_text is None: logger.warning("TABLE_MODIFY_CELL: 'table_coords' или 'new_cell_text' не указаны."); return False
    row_idx, col_idx = table_coords.get("row"), table_coords.get("col")
    if row_idx is None or col_idx is None: logger.warning("TABLE_MODIFY_CELL: 'row' или 'col' в coords не указаны."); return False

    table = get_table_by_description(doc, target_description)
    if not table: logger.warning(f"TABLE_MODIFY_CELL: Таблица '{target_description}' не найдена."); return False
    
    try:
        cell_to_modify: CellType = table.cell(row_idx, col_idx)
        while len(cell_to_modify.paragraphs) > 1: p_el = cell_to_modify.paragraphs[-1]._element; p_el.getparent().remove(p_el)
        first_para = cell_to_modify.paragraphs[0] if cell_to_modify.paragraphs else cell_to_modify.add_paragraph()
        first_para.text = new_cell_text
        logger.info(f"TABLE_MODIFY_CELL: Ячейка ({row_idx},{col_idx}) изменена на '{new_cell_text}'.")
        return True
    except IndexError: logger.warning(f"TABLE_MODIFY_CELL: Индекс ({row_idx},{col_idx}) вне диапазона."); return False
    except Exception as e: logger.error(f"TABLE_MODIFY_CELL: Ошибка: {e}"); return False

def handle_table_add_row(doc: Document, target_description: dict, parameters: dict) -> bool:
    # ... (ваш существующий код _handle_table_add_row, использующий get_table_by_description) ...
    logger.info(f"Выполнение TABLE_ADD_ROW: target={target_description}, params={parameters}")
    row_data = parameters.get("row_data")
    insert_at_index = parameters.get("insert_at_index")

    if not isinstance(row_data, list): logger.warning("TABLE_ADD_ROW: 'row_data' не список."); return False
    table = get_table_by_description(doc, target_description)
    if not table: logger.warning(f"TABLE_ADD_ROW: Таблица '{target_description}' не найдена."); return False
    if len(row_data) != len(table.columns): logger.warning(f"TABLE_ADD_ROW: Данные не совпадают с кол-вом колонок."); return False

    # Логика вставки по индексу (упрощенная - всегда в конец)
    if insert_at_index is not None: logger.warning("TABLE_ADD_ROW: Вставка по индексу пока не полностью поддерживается, строка добавлена в конец.")
    new_row = table.add_row()
    for i, cell_text in enumerate(row_data): new_row.cells[i].text = str(cell_text)
    logger.info(f"TABLE_ADD_ROW: Строка {row_data} добавлена в таблицу.")
    return True