# core/operations/text_operations.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX # WD_COLOR_INDEX для подсветки
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from loguru import logger
from ..docx_utils import find_paragraphs_with_text, find_runs_with_text # Используем относительный импорт

# Ваша функция _replace_text_in_paragraph_runs_with_highlight должна быть здесь
# или в docx_utils.py, если она достаточно общая.
# Давайте предположим, что она остается здесь, так как тесно связана с заменой.
def _replace_text_in_paragraph_runs_with_highlight(p: Paragraph, old_text: str, new_text: str) -> bool:
    """
    Находит и заменяет текст в абзаце, который может быть разбит на несколько 'runs'.
    Версия без подсветки.
    """
    if not old_text or old_text not in p.text:
        return False

    runs = p.runs
    
    # --- Простая замена (если весь old_text находится в одном run) ---
    for i, run in enumerate(runs):
        if old_text in run.text:
            # Проверяем, что это не просто часть совпадения, разбросанного по нескольким run'ам
            run_text_start_in_para = sum(len(r.text) for r in runs[:i])
            para_find_start = p.text.find(old_text)
            
            # Если начальная позиция совпадения в полном тексте абзаца находится внутри текущего run'а
            if para_find_start >= run_text_start_in_para and \
               para_find_start + len(old_text) <= run_text_start_in_para + len(run.text):
                
                logger.debug(f" (простая замена): old_text ('{old_text}') найден и заменяется в одном run: '{run.text}'")
                current_run_text = run.text
                start_replace_index = current_run_text.find(old_text)
                
                if start_replace_index != -1:
                    end_replace_index = start_replace_index + len(old_text)
                    run.text = current_run_text[:start_replace_index] + new_text + current_run_text[end_replace_index:]
                    # СТРОКА С ПОДСВЕТКОЙ УДАЛЕНА
                    return True
    
    # --- Сложная замена (если old_text разбит на несколько run'ов) ---
    current_pos = 0
    text_segments = []
    para_full_text = ""
    for i, run_obj_iter in enumerate(runs):
        text_segments.append({'index': i, 'start_pos': current_pos, 'text': run_obj_iter.text, 'obj': run_obj_iter})
        para_full_text += run_obj_iter.text
        current_pos += len(run_obj_iter.text)
        
    start_match_idx = para_full_text.find(old_text)
    if start_match_idx != -1:
        end_match_idx = start_match_idx + len(old_text)
        first_run_involved_details = None
        
        for seg in text_segments:
            if seg['start_pos'] <= start_match_idx < (seg['start_pos'] + len(seg['text'])):
                first_run_involved_details = seg
                break
        
        if first_run_involved_details:
            first_run_obj = first_run_involved_details['obj']
            offset_in_first_run = start_match_idx - first_run_involved_details['start_pos']
            prefix = first_run_obj.text[:offset_in_first_run]
            len_old_text_in_first_run_after_prefix = len(first_run_involved_details['text']) - offset_in_first_run
            
            first_run_obj.text = prefix + new_text
            # СТРОКА С ПОДСВЕТКОЙ УДАЛЕНА
            
            remaining_old_text_to_remove_len = len(old_text) - len_old_text_in_first_run_after_prefix
            if remaining_old_text_to_remove_len < 0:
                remaining_old_text_to_remove_len = 0
            
            for k in range(first_run_involved_details['index'] + 1, len(text_segments)):
                if remaining_old_text_to_remove_len <= 0:
                    break
                current_run_seg = text_segments[k]
                current_run_obj = current_run_seg['obj']
                if len(current_run_obj.text) <= remaining_old_text_to_remove_len:
                    remaining_old_text_to_remove_len -= len(current_run_obj.text)
                    current_run_obj.text = ""
                else:
                    current_run_obj.text = current_run_obj.text[remaining_old_text_to_remove_len:]
                    remaining_old_text_to_remove_len = 0
            
            return True
            
    return False


def handle_replace_text(doc: Document, target_description: dict, parameters: dict) -> bool:
    # ... (ваш существующий код _handle_replace_text, но использующий _replace_text_in_paragraph_runs_with_highlight отсюда
    # и find_paragraphs_with_text из docx_utils) ...
    logger.info(f"Выполнение REPLACE_TEXT: target={target_description}, params={parameters}")
    old_text = parameters.get("old_text")
    new_text = parameters.get("new_text", "")
    context_text = target_description.get("text_to_find")
    placeholder = target_description.get("placeholder")

    if placeholder and not old_text:
        old_text = placeholder
        logger.info(f"Используем плейсхолдер '{placeholder}' как old_text.")
    if not old_text:
        logger.warning("REPLACE_TEXT: 'old_text' или 'placeholder' не указан или пуст.")
        return False

    modified_count = 0
    elements_to_search_in = []
    if context_text:
        elements_to_search_in.extend(find_paragraphs_with_text(doc, context_text))
        for section in doc.sections:
            elements_to_search_in.extend(find_paragraphs_with_text(section.header, context_text))
            elements_to_search_in.extend(find_paragraphs_with_text(section.footer, context_text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    elements_to_search_in.extend(find_paragraphs_with_text(cell, context_text))
        if not elements_to_search_in: logger.warning(f"REPLACE_TEXT: Контекстный текст '{context_text}' не найден.")

    if not elements_to_search_in:
        elements_to_search_in.extend(doc.paragraphs)
        for section in doc.sections:
            elements_to_search_in.extend(section.header.paragraphs)
            elements_to_search_in.extend(section.footer.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    elements_to_search_in.extend(cell.paragraphs)
    
    unique_paragraphs = list(dict.fromkeys(elements_to_search_in))
    for p in unique_paragraphs:
        if _replace_text_in_paragraph_runs_with_highlight(p, old_text, new_text):
            modified_count += 1
    
    if modified_count > 0: logger.info(f"REPLACE_TEXT: Текст '{old_text}' заменен на '{new_text}' в {modified_count} местах."); return True
    else: logger.warning(f"REPLACE_TEXT: Текст '{old_text}' не найден для замены."); return False


def handle_insert_text(doc: Document, target_description: dict, parameters: dict) -> bool:
    logger.info(f"Выполнение INSERT_TEXT: target={target_description}, params={parameters}")
    text_to_insert = parameters.get("text_to_insert")
    position = parameters.get("position")
    target_text = target_description.get("text_to_find")

    if not all([text_to_insert, position, target_text]):
        logger.warning("INSERT_TEXT: Не все параметры указаны.")
        return False
    
    target_paragraphs = find_paragraphs_with_text(doc, target_text)
    if not target_paragraphs:
        logger.warning(f"INSERT_TEXT: Абзац с текстом '{target_text}' не найден.")
        return False
    
    target_p = target_paragraphs[0]
    if len(target_paragraphs) > 1:
        logger.warning(f"INSERT_TEXT: Найдено несколько абзацев с '{target_text}'. Используется первый.")

    # --- ЛОГИКА СОХРАНЕНИЯ СТИЛЯ ---

    if position == "after_paragraph":
        # Создаем новый абзац и ПРИМЕНЯЕМ К НЕМУ СТИЛЬ целевого абзаца
        new_p = target_p.insert_paragraph_before('') # Создаем пустой абзац перед целевым
        new_p.style = target_p.style # Копируем стиль абзаца (выравнивание, отступы и т.д.)
        new_p.add_run(text_to_insert).bold = target_p.runs[-1].bold if target_p.runs else None
        # Копируем стиль последнего run-а (жирный, курсив и т.д.)
        if target_p.runs:
            last_run_style = target_p.runs[-1].font
            new_run = new_p.runs[0]
            new_run.font.name = last_run_style.name
            new_run.font.size = last_run_style.size
            new_run.font.bold = last_run_style.bold
            new_run.font.italic = last_run_style.italic
            new_run.font.underline = last_run_style.underline
            new_run.font.color.rgb = last_run_style.color.rgb
        
        # Перемещаем созданный абзац ПОСЛЕ целевого
        p_element = target_p._element
        new_p_element = new_p._element
        p_element.addnext(new_p_element)

    elif position == "before_paragraph":
        # Метод insert_paragraph_before уже хорошо справляется с копированием стиля абзаца.
        # Мы просто добавим копирование стиля текста (run).
        new_p = target_p.insert_paragraph_before(text_to_insert, style=target_p.style)
        if target_p.runs:
            # Применим стиль первого run'а целевого абзаца к нашему новому run'у
            first_run_style = target_p.runs[0].font
            new_run = new_p.runs[0]
            new_run.font.name = first_run_style.name
            new_run.font.size = first_run_style.size
            new_run.font.bold = first_run_style.bold
            new_run.font.italic = first_run_style.italic
            new_run.font.underline = first_run_style.underline
            new_run.font.color.rgb = first_run_style.color.rgb

    elif position == "start_of_paragraph":
        # Добавляем run в начало, копируя стиль первого существующего run'а
        # Важно: Сначала добавляем пробел, чтобы текст не слипся
        text_to_insert_with_space = text_to_insert + " "
        first_run = target_p.runs[0] if target_p.runs else None
        
        # Создаем новый run и вставляем его в начало
        new_run = target_p.add_run() # Временный пустой run в конце
        new_run._element.getparent().remove(new_run._element) # Отсоединяем его
        target_p._p.insert(0, new_run._element) # Вставляем в начало
        new_run.text = text_to_insert_with_space
        
        # Копируем стиль
        if first_run:
            new_run.style = first_run.style # Копируем стиль символов
            # Явно копируем форматирование шрифта
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.font.bold = first_run.font.bold
            new_run.font.italic = first_run.font.italic
            new_run.font.underline = first_run.font.underline
            new_run.font.color.rgb = first_run.font.color.rgb
        else: # Если абзац был пустой
             new_run.style = target_p.style

    elif position == "end_of_paragraph":
        # Добавляем run в конец, копируя стиль последнего существующего run'а
        # Важно: Сначала добавляем пробел, чтобы текст не слипся
        text_to_insert_with_space = " " + text_to_insert
        last_run = target_p.runs[-1] if target_p.runs else None
        
        new_run = target_p.add_run(text_to_insert_with_space)
        
        # Копируем стиль
        if last_run:
            new_run.style = last_run.style # Копируем стиль символов
            # Явно копируем форматирование шрифта
            new_run.font.name = last_run.font.name
            new_run.font.size = last_run.font.size
            new_run.font.bold = last_run.font.bold
            new_run.font.italic = last_run.font.italic
            new_run.font.underline = last_run.font.underline
            new_run.font.color.rgb = last_run.font.color.rgb
        else: # Если абзац был пустой
            new_run.style = target_p.style

    else:
        logger.warning(f"INSERT_TEXT: Неизвестная позиция '{position}'.")
        return False
        
    logger.info(f"INSERT_TEXT: Текст '{text_to_insert}' вставлен {position} относительно '{target_text}' с сохранением стиля.")
    return True

def _apply_single_formatting_rule_to_run(run, rule: dict): # Отдельная для Run
    style = rule.get("style"); value = rule.get("value")
    font = run.font
    if style == "bold": font.bold = bool(value)
    elif style == "italic": font.italic = bool(value)
    # ... (остальные стили шрифта из вашего _apply_single_formatting_rule_to_element) ...
    elif style == "underline": font.underline = bool(value)
    elif style == "font_size" and isinstance(value, (int, float)): font.size = Pt(value)
    elif style == "font_name": font.name = str(value)
    elif style == "font_color_rgb":
        try: font.color.rgb = RGBColor.from_string(str(value).replace("#",""))
        except Exception as e: logger.warning(f"Неверный RGB цвет '{value}': {e}")
    elif style == "highlight_color":
        color_val = str(value).upper()
        if hasattr(WD_COLOR_INDEX, color_val): font.highlight_color = getattr(WD_COLOR_INDEX, color_val)
        elif color_val == "NONE": font.highlight_color = None
        else: logger.warning(f"Неизвестный цвет выделения '{value}'")


def _copy_run_style(source_run, target_run):
    """Копирует все атрибуты форматирования с одного run'а на другой."""
    target_run.style = source_run.style
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.highlight_color = source_run.font.highlight_color
    # Добавьте другие атрибуты, если они вам нужны (например, sub/superscript)


def _format_text_within_paragraph(p, text_to_format, rules):
    """
    Находит текст в абзаце и форматирует ТОЛЬКО его, разбивая run'ы при необходимости.
    Эта версия более простая и надежная.
    """
    if text_to_format not in p.text:
        return False

    # Сохраняем копии оригинальных run'ов, чтобы итерироваться по ним
    runs = list(p.runs)
    # Очищаем абзац, чтобы пересобрать его
    for run in runs:
        p._p.remove(run._r)
        
    full_text = "".join(r.text for r in runs)
    current_pos_in_text = 0
    run_idx = 0

    # Проходим по всему тексту оригинального абзаца
    while current_pos_in_text < len(full_text):
        original_run = runs[run_idx]
        run_text = original_run.text
        
        # Ищем совпадение ВНУТРИ текущего run'а
        pos_in_run = run_text.find(text_to_format)
        
        if pos_in_run != -1:
            # Нашли! Разбиваем run на три части
            before_text = run_text[:pos_in_run]
            after_text = run_text[pos_in_run + len(text_to_format):]

            # 1. Часть "до"
            if before_text:
                new_run = p.add_run(before_text)
                _copy_run_style(original_run, new_run)
            
            # 2. Форматируемая часть
            formatted_run = p.add_run(text_to_format)
            _copy_run_style(original_run, formatted_run) # Сначала копируем старый стиль
            for rule in rules: # Потом применяем новый поверх
                _apply_single_formatting_rule_to_run(formatted_run, rule)

            # 3. Часть "после"
            if after_text:
                # Создаем временный run, чтобы продолжить обработку остатка
                original_run.text = after_text 
                run_text = after_text
                # Ищем совпадение в остатке этого же run'а
                continue 
            
        else:
            # Совпадений в этом run'е нет, просто копируем его как есть
            new_run = p.add_run(run_text)
            _copy_run_style(original_run, new_run)

        # Переходим к следующему оригинальному run'у
        current_pos_in_text += len(run_text)
        run_idx += 1
        if run_idx >= len(runs):
            break
            
    return True


def handle_apply_text_formatting(doc: Document, target_description: dict, parameters: dict) -> bool:
    """
    Применяет форматирование к конкретным сегментам текста.
    """
    logger.info(f"Выполнение APPLY_TEXT_FORMATTING: target={target_description}, params={parameters}")
    
    target_text_context = target_description.get("text_to_find")
    apply_to_text_segment = parameters.get("apply_to_text_segment")
    formatting_rules = parameters.get("formatting_rules", [])

    if not formatting_rules or not apply_to_text_segment:
        logger.warning("APPLY_TEXT_FORMATTING: 'formatting_rules' или 'apply_to_text_segment' не указаны.")
        return False
        
    # ИЗМЕНЕНИЕ: Ищем абзацы, содержащие точный сегмент, а не контекст.
    # Это более надежно, если LLM ошиблась с контекстом.
    paragraphs_to_process = find_paragraphs_with_text(doc, apply_to_text_segment, partial_match=True)
    if not paragraphs_to_process:
        logger.warning(f"APPLY_TEXT_FORMATTING: Сегмент для форматирования '{apply_to_text_segment}' не найден в документе.")
        return False
    
    modified_something = False
    for p in paragraphs_to_process:
        # Внутри найденного абзаца применяем форматирование
        if _format_text_within_paragraph(p, apply_to_text_segment, formatting_rules):
            modified_something = True
    
    if modified_something:
        logger.info(f"APPLY_TEXT_FORMATTING: Форматирование для сегмента '{apply_to_text_segment}' применено.")

    return modified_something