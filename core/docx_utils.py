# core/docx_utils.py
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from loguru import logger
from docx.section import _Header, _Footer
from typing import Union, List
import sys

ContainerType = Union[Document, _Cell, _Header, _Footer, Paragraph]

def find_paragraphs_with_text(container: ContainerType, 
                              text_to_find: str, 
                              partial_match: bool = False) -> List[Paragraph]:
    """
    Находит все абзацы в заданном контейнере, содержащие указанный текст.

    Args:
        container: Объект, в котором искать (Document, _Cell, Section.header/footer и т.д.).
        text_to_find (str): Текст для поиска.
        partial_match (bool): Если True, ищет вхождение текста (text_to_find in p.text).
                              Если False (по умолчанию), ищет точное совпадение текста абзаца (p.text == text_to_find).

    Returns:
        list[Paragraph]: Список найденных объектов абзацев.
    """
    if container is None:
        logger.trace(f"find_paragraphs_with_text: Контейнер is None, поиск текста '{text_to_find}' невозможен.")
        return []
    if not hasattr(container, 'paragraphs'):
        logger.trace(f"find_paragraphs_with_text: Контейнер типа {type(container)} не имеет атрибута 'paragraphs'. Поиск текста '{text_to_find}' невозможен.")
        return []
    if not text_to_find: # Проверка на пустой text_to_find
        logger.trace(f"find_paragraphs_with_text: text_to_find пуст, поиск не выполняется в контейнере {type(container)}.")
        return []

    logger.trace(f"find_paragraphs_with_text: Поиск '{text_to_find}' в контейнере типа {type(container)}, partial_match={partial_match}")
    
    found_paragraphs = []
    try:
        for p_idx, p in enumerate(container.paragraphs):
            p_text_stripped = p.text.strip() # Часто полезно для сравнения без крайних пробелов
            text_to_find_stripped = text_to_find.strip()

            logger.trace(f"  Проверка абзаца #{p_idx}: repr(p.text)='{repr(p.text)}'")
            logger.trace(f"    Искомый текст (repr): '{repr(text_to_find)}'")
            logger.trace(f"    p.text.strip() (repr): '{repr(p_text_stripped)}'")
            logger.trace(f"    text_to_find.strip() (repr): '{repr(text_to_find_stripped)}'")

            match_found = False
            if partial_match:
                if text_to_find in p.text:
                    logger.debug(f"    ЧАСТИЧНОЕ СОВПАДЕНИЕ НАЙДЕНО: '{text_to_find}' in '{p.text}'")
                    match_found = True
                # Дополнительная проверка для частичного совпадения без учета крайних пробелов
                elif text_to_find_stripped and p_text_stripped and text_to_find_stripped in p_text_stripped:
                    logger.debug(f"    ЧАСТИЧНОЕ СОВПАДЕНИЕ (со strip) НАЙДЕНО: '{text_to_find_stripped}' in '{p_text_stripped}'")
                    match_found = True
            else: # Точное совпадение
                if p.text == text_to_find:
                    logger.debug(f"    ТОЧНОЕ СОВПАДЕНИЕ НАЙДЕНО: p.text == '{text_to_find}'")
                    match_found = True
                # Дополнительная проверка для точного совпадения без учета крайних пробелов
                elif p_text_stripped == text_to_find_stripped:
                    logger.debug(f"    ТОЧНОЕ СОВПАДЕНИЕ (со strip) НАЙДЕНО: p.text.strip() == '{text_to_find_stripped}'")
                    match_found = True
            
            if match_found:
                found_paragraphs.append(p)
            else:
                logger.trace(f"    Совпадение не найдено для абзаца #{p_idx}.")
                
    except Exception as e:
        logger.error(f"find_paragraphs_with_text: Ошибка при итерации по абзацам в контейнере {type(container)}: {e}")

    if found_paragraphs:
        logger.debug(f"find_paragraphs_with_text: Найдено {len(found_paragraphs)} абзац(ев) с текстом (или его частью) '{text_to_find}'.")
    else:
        logger.trace(f"find_paragraphs_with_text: Абзацы с текстом (или его частью) '{text_to_find}' не найдены в контейнере {type(container)}.")
        
    return found_paragraphs

def find_runs_with_text(paragraph: Paragraph, text_to_find: str) -> list: # list of Run objects
    """Находит все runs в абзаце, содержащие text_to_find."""
    found_runs = []
    if text_to_find:
        for run in paragraph.runs:
            if text_to_find in run.text:
                found_runs.append(run)
    return found_runs

def get_table_by_description(doc: Document, target_description: dict) -> Table | None:
    """Находит таблицу по описанию (индекс или текст)."""
    table_index = target_description.get("table_index")
    text_to_find = target_description.get("text_to_find")

    if table_index is not None:
        if 0 <= table_index < len(doc.tables):
            return doc.tables[table_index]
        else:
            logger.warning(f"Индекс таблицы {table_index} вне диапазона.")
            return None

    if text_to_find:
        for i, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    if text_to_find in cell.text:
                        logger.info(f"Таблица найдена по тексту '{text_to_find}' в ячейке (индекс {i}).")
                        return table
        logger.warning(f"Таблица с текстом '{text_to_find}' не найдена внутри ячеек.")
    
    if not doc.tables:
        logger.warning("В документе нет таблиц.")
        return None
        
    logger.warning(f"Не удалось однозначно идентифицировать таблицу по описанию: {target_description}")
    return None

def extract_text_from_doc(doc_object: Document) -> str:
    """Извлекает весь видимый текст из документа для передачи в LLM."""
    # ... (ваш существующий код extract_text_from_doc) ...
    full_text_parts = []
    for p in doc_object.paragraphs:
        full_text_parts.append(p.text)
    # ... и так далее для таблиц, колонтитулов ...
    for table in doc_object.tables:
         for row in table.rows:
             for cell in row.cells:
                 for p_in_cell in cell.paragraphs:
                     full_text_parts.append(p_in_cell.text)
    for section in doc_object.sections:
         for p_in_header in section.header.paragraphs:
             full_text_parts.append(p_in_header.text)
         for table_in_header in section.header.tables: # Добавлено
             for row in table_in_header.rows:
                 for cell in row.cells:
                     for p_in_cell in cell.paragraphs:
                         full_text_parts.append(p_in_cell.text)
         for p_in_footer in section.footer.paragraphs:
             full_text_parts.append(p_in_footer.text)
         for table_in_footer in section.footer.tables: # Добавлено
             for row in table_in_footer.rows:
                 for cell in row.cells:
                     for p_in_cell in cell.paragraphs:
                         full_text_parts.append(p_in_cell.text)
    return "\n".join(full_text_parts)